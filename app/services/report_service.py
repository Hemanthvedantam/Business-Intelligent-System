# Report Service — generates PDF and DOCX investigation reports.
#
# Upgrade 7: produces downloadable reports containing:
#   Summary · Charts · Evidence Tree · Root Causes · Recommendations · Forecast
#
# PDF:  uses reportlab (pure Python, no wkhtmltopdf dependency)
# DOCX: uses python-docx

import io
import re
import os
from datetime import datetime
from typing import Optional

from app.core.logging import get_logger

logger = get_logger(__name__)


# ─────────────────────────────────────────────────────────────────────────────
#  Public API
# ─────────────────────────────────────────────────────────────────────────────

def generate_pdf(investigation: object, result: Optional[dict] = None) -> bytes:
    """
    Generate a PDF report for the given investigation.
    `investigation` is a SQLAlchemy Investigation model instance.
    `result` is the optional pre-built result dict; if omitted it is
    reconstructed from the model fields.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer,
            Table, TableStyle, HRFlowable, ListFlowable, ListItem,
        )
    except ImportError:
        raise RuntimeError(
            "reportlab is required for PDF export. "
            "Run: pip install reportlab --break-system-packages"
        )

    data = _build_report_data(investigation, result)
    buf  = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=f"ABIP Investigation Report — {data['question'][:60]}",
    )

    styles  = getSampleStyleSheet()
    PRIMARY = colors.HexColor("#2563EB")
    DARK    = colors.HexColor("#0F172A")
    MUTED   = colors.HexColor("#64748B")
    RED     = colors.HexColor("#EF4444")
    GREEN   = colors.HexColor("#22C55E")
    BG      = colors.HexColor("#F8FAFC")
    BORDER  = colors.HexColor("#E2E8F0")

    h1 = ParagraphStyle("H1", parent=styles["Heading1"],
                         fontSize=20, textColor=DARK, spaceAfter=4,
                         fontName="Helvetica-Bold")
    h2 = ParagraphStyle("H2", parent=styles["Heading2"],
                         fontSize=13, textColor=PRIMARY, spaceBefore=14,
                         spaceAfter=4, fontName="Helvetica-Bold")
    body = ParagraphStyle("Body", parent=styles["Normal"],
                          fontSize=10, textColor=DARK, spaceAfter=6,
                          leading=15)
    muted = ParagraphStyle("Muted", parent=styles["Normal"],
                           fontSize=9, textColor=MUTED, spaceAfter=4)
    mono  = ParagraphStyle("Mono", parent=styles["Code"],
                           fontSize=8.5, textColor=DARK, spaceAfter=4,
                           fontName="Courier")

    story = []

    # ── Cover / Header ────────────────────────────────────────────────────
    story.append(Paragraph("ABIP — Investigation Report", h1))
    story.append(HRFlowable(width="100%", thickness=1, color=BORDER))
    story.append(Spacer(1, 4 * mm))

    meta_rows = [
        ["Question",  data["question"]],
        ["Dataset",   data["filename"]],
        ["Domain",    data["domain"].capitalize()],
        ["Generated", data["generated_at"]],
        ["Status",    data["status"]],
    ]
    if data.get("confidence"):
        meta_rows.append(["Confidence", f"{data['confidence']}%"])

    meta_table = Table(
        meta_rows,
        colWidths=[35 * mm, None],
        hAlign="LEFT",
    )
    meta_table.setStyle(TableStyle([
        ("FONTNAME",    (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE",    (0, 0), (-1, -1), 9),
        ("TEXTCOLOR",   (0, 0), (0, -1), MUTED),
        ("TEXTCOLOR",   (1, 0), (1, -1), DARK),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING",    (0, 0), (-1, -1), 2),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 6 * mm))

    # ── Executive Summary ─────────────────────────────────────────────────
    story.append(Paragraph("Executive Summary", h2))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
    story.append(Spacer(1, 2 * mm))

    if data.get("final_summary"):
        for para in data["final_summary"].strip().split("\n\n"):
            story.append(Paragraph(_clean(para), body))

    # KPI metrics table
    metrics = data.get("metrics", [])
    if metrics:
        story.append(Spacer(1, 3 * mm))
        metric_rows = [[m["label"], m["value"]] for m in metrics]
        mt = Table(metric_rows, colWidths=[60 * mm, 80 * mm], hAlign="LEFT")
        mt.setStyle(TableStyle([
            ("BACKGROUND",  (0, 0), (-1, -1), BG),
            ("FONTNAME",    (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTSIZE",    (0, 0), (-1, -1), 9),
            ("TEXTCOLOR",   (0, 0), (0, -1), MUTED),
            ("TEXTCOLOR",   (1, 0), (1, -1), DARK),
            ("BOX",         (0, 0), (-1, -1), 0.5, BORDER),
            ("INNERGRID",   (0, 0), (-1, -1), 0.25, BORDER),
            ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.white, BG]),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING",    (0, 0), (-1, -1), 5),
            ("LEFTPADDING",   (0, 0), (-1, -1), 8),
        ]))
        story.append(mt)

    # ── Evidence Tree ─────────────────────────────────────────────────────
    evidence_tree = data.get("evidence_tree")
    if evidence_tree and evidence_tree.get("branches"):
        story.append(Paragraph("Evidence Tree", h2))
        story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
        story.append(Spacer(1, 2 * mm))
        story.append(Paragraph(f"▶  {evidence_tree['root']}", body))

        branches = evidence_tree["branches"]
        for i, b in enumerate(branches):
            prefix = "└─" if i == len(branches) - 1 else "├─"
            val    = f"  {b['value']}" if b.get("value") else ""
            story.append(
                Paragraph(
                    f"<font color='#{_dir_hex(b.get('direction', 'neutral'))}'>"
                    f"{prefix} {b['label']}{val}</font>",
                    mono,
                )
            )

    # ── Root Causes ───────────────────────────────────────────────────────
    causes = data.get("root_causes", [])
    if causes:
        story.append(Paragraph("Root Causes", h2))
        story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
        story.append(Spacer(1, 2 * mm))
        items = [
            ListItem(Paragraph(_clean(c), body), bulletColor=RED, leftIndent=12)
            for c in causes
        ]
        story.append(ListFlowable(items, bulletType="bullet", start="•"))

    # ── Recommendations ───────────────────────────────────────────────────
    recs = data.get("recommendations", [])
    if recs:
        story.append(Paragraph("Recommendations", h2))
        story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
        story.append(Spacer(1, 2 * mm))
        items = [
            ListItem(Paragraph(f"{i + 1}.  {_clean(r)}", body),
                     bulletColor=GREEN, leftIndent=12)
            for i, r in enumerate(recs)
        ]
        story.append(ListFlowable(items, bulletType="bullet", start="•"))

    # ── Forecast ─────────────────────────────────────────────────────────
    forecast = data.get("forecast")
    if forecast and isinstance(forecast, dict):
        story.append(Paragraph("Forecast", h2))
        story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
        story.append(Spacer(1, 2 * mm))
        preds = forecast.get("predictions") or forecast.get("summary", "")
        if preds:
            story.append(Paragraph(_clean(str(preds)), body))

    # ── Footer ────────────────────────────────────────────────────────────
    story.append(Spacer(1, 10 * mm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BORDER))
    story.append(Paragraph(
        f"Generated by ABIP — AI Business Intelligence Platform · {data['generated_at']}",
        muted,
    ))

    doc.build(story)
    return buf.getvalue()


def generate_docx(investigation: object, result: Optional[dict] = None) -> bytes:
    """
    Generate a DOCX report for the given investigation.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError(
            "python-docx is required for DOCX export. "
            "Run: pip install python-docx --break-system-packages"
        )

    data = _build_report_data(investigation, result)
    doc  = Document()

    # ── Page margins ──────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    def heading(text: str, level: int = 1, color=(37, 99, 235)):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
        return p

    def para(text: str, bold: bool = False, size: int = 10):
        p = doc.add_paragraph()
        run = p.add_run(_clean(text))
        run.font.size = Pt(size)
        run.bold = bold
        return p

    def add_hr():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "E2E8F0")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Cover ─────────────────────────────────────────────────────────────
    t = doc.add_heading("ABIP — Investigation Report", 0)
    for run in t.runs:
        run.font.color.rgb = RGBColor(15, 23, 42)
    add_hr()

    # Meta table
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in [
        ("Question",  data["question"]),
        ("Dataset",   data["filename"]),
        ("Domain",    data["domain"].capitalize()),
        ("Generated", data["generated_at"]),
        ("Status",    data["status"]),
    ] + ([("Confidence", f"{data['confidence']}%")] if data.get("confidence") else []):
        row = table.add_row().cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = str(value)

    doc.add_paragraph()

    # ── Executive Summary ─────────────────────────────────────────────────
    heading("Executive Summary", level=1)
    add_hr()

    if data.get("final_summary"):
        for p in data["final_summary"].strip().split("\n\n"):
            para(p)

    # Metrics
    metrics = data.get("metrics", [])
    if metrics:
        mt = doc.add_table(rows=1, cols=2)
        mt.style = "Table Grid"
        mt.rows[0].cells[0].text = "Metric"
        mt.rows[0].cells[1].text = "Value"
        for m in metrics:
            row = mt.add_row().cells
            row[0].text = m["label"]
            row[1].text = m["value"]

    doc.add_paragraph()

    # ── Evidence Tree ─────────────────────────────────────────────────────
    evidence_tree = data.get("evidence_tree")
    if evidence_tree and evidence_tree.get("branches"):
        heading("Evidence Tree", level=1)
        add_hr()
        para(f"▶  {evidence_tree['root']}", bold=True)
        branches = evidence_tree["branches"]
        for i, b in enumerate(branches):
            prefix = "└─" if i == len(branches) - 1 else "├─"
            val    = f"  {b['value']}" if b.get("value") else ""
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{prefix} {b['label']}{val}")
            run.font.size = Pt(10)
        doc.add_paragraph()

    # ── Root Causes ───────────────────────────────────────────────────────
    causes = data.get("root_causes", [])
    if causes:
        heading("Root Causes", level=1)
        add_hr()
        for c in causes:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(_clean(c)).font.size = Pt(10)
        doc.add_paragraph()

    # ── Recommendations ───────────────────────────────────────────────────
    recs = data.get("recommendations", [])
    if recs:
        heading("Recommendations", level=1)
        add_hr()
        for i, r in enumerate(recs):
            p = doc.add_paragraph(style="List Number")
            p.add_run(_clean(r)).font.size = Pt(10)
        doc.add_paragraph()

    # ── Forecast ──────────────────────────────────────────────────────────
    forecast = data.get("forecast")
    if forecast and isinstance(forecast, dict):
        heading("Forecast", level=1)
        add_hr()
        preds = forecast.get("predictions") or forecast.get("summary", "")
        if preds:
            para(str(preds))

    # ── Footer ────────────────────────────────────────────────────────────
    footer_para = doc.sections[0].footer.paragraphs[0]
    footer_para.text = f"Generated by ABIP — AI Business Intelligence Platform · {data['generated_at']}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
#  Internal helpers
# ─────────────────────────────────────────────────────────────────────────────

def _build_report_data(investigation, result: Optional[dict]) -> dict:
    """
    Merge model fields + optional result dict into one flat data dict
    that both generate_pdf and generate_docx consume.
    """
    if result is None:
        result = {}

    # Pull from executive_summary if present
    exec_summary = result.get("executive_summary", {}) or {}
    metrics      = exec_summary.get("metrics", [])
    confidence   = exec_summary.get("confidence") or result.get("confidence")
    evidence_tree = result.get("evidence_tree") or {}

    # Normalise root_causes / recommendations to plain strings
    root_causes    = _normalise_list(result.get("root_causes")    or getattr(investigation, "root_causes",    []) or [])
    recommendations= _normalise_list(result.get("recommendations") or getattr(investigation, "recommendations", []) or [])

    return {
        "question":      getattr(investigation, "question",          "Investigation"),
        "filename":      getattr(investigation, "dataset_filename",  "Unknown"),
        "domain":        getattr(investigation, "domain",            "other") or "other",
        "status":        str(getattr(investigation, "status",        "completed")),
        "generated_at":  datetime.now().strftime("%d %b %Y  %H:%M"),
        "confidence":    confidence,
        "final_summary": result.get("final_summary") or getattr(investigation, "final_summary", ""),
        "metrics":       metrics,
        "evidence_tree": evidence_tree,
        "root_causes":   root_causes,
        "recommendations": recommendations,
        "forecast":      result.get("forecast") or {},
    }


def _normalise_list(items) -> list:
    """Convert a list of anything to a list of plain strings."""
    out = []
    for item in (items or []):
        if isinstance(item, str):
            out.append(item.strip())
        elif isinstance(item, dict):
            # e.g. {"title": "...", "detail": "..."}
            title  = item.get("title",  "")
            detail = item.get("detail", "")
            out.append(f"{title}: {detail}" if title else detail)
        else:
            out.append(str(item))
    return [i for i in out if i]


def _clean(text: str) -> str:
    """Strip markdown bold/italic/code markers."""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*",     r"\1", text)
    text = re.sub(r"`(.*?)`",       r"\1", text)
    return text.strip()


def _dir_hex(direction: str) -> str:
    """Return a hex colour string (without #) for evidence tree directions."""
    return {"negative": "EF4444", "positive": "22C55E"}.get(direction, "64748B")